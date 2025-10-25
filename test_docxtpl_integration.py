#!/usr/bin/env python3
"""Test docxtpl integration with both simple and advanced features."""

import tempfile
from pathlib import Path
from docx import Document
from docx_processor import DocxProcessor


def create_simple_template(path):
    """Create a simple test template with {{placeholders}}."""
    doc = Document()
    doc.add_heading('測試文件', 0)
    doc.add_paragraph('姓名：{{name}}')
    doc.add_paragraph('日期：{{date}}')
    doc.add_paragraph('公司：{{company.name}}')
    doc.add_paragraph('地址：{{company.address}}')
    doc.save(path)
    print(f"✓ Created simple template: {path}")


def create_advanced_template(path):
    """Create an advanced template with Jinja2 syntax."""
    doc = Document()
    doc.add_heading('進階測試文件', 0)
    doc.add_paragraph('員工：{{name}}')
    doc.add_paragraph('')
    doc.add_paragraph('專案列表：')
    doc.add_paragraph('{% for project in projects %}')
    doc.add_paragraph('  - {{project.name}}: {{project.budget}}元')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('')
    doc.add_paragraph('{% if salary > 40000 %}')
    doc.add_paragraph('薪資等級：高')
    doc.add_paragraph('{% else %}')
    doc.add_paragraph('薪資等級：一般')
    doc.add_paragraph('{% endif %}')
    doc.save(path)
    print(f"✓ Created advanced template: {path}")


def test_simple_replacement():
    """Test simple {{key}} replacement."""
    print("\n=== Test 1: Simple Replacement ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "simple_template.docx"
        output_path = Path(tmpdir) / "simple_output.docx"

        # Create template
        create_simple_template(template_path)

        # Prepare data
        data = {
            'name': '張三',
            'date': '2024-10-23',
            'company': {
                'name': 'ABC公司',
                'address': '台北市信義區'
            }
        }

        # Process
        processor = DocxProcessor(template_path)
        processor.fill_template(data, output_path, verbose=True)

        # Verify
        doc = Document(output_path)
        text = '\n'.join([p.text for p in doc.paragraphs])

        assert '張三' in text, "Name replacement failed"
        assert '2024-10-23' in text, "Date replacement failed"
        assert 'ABC公司' in text, "Nested company.name replacement failed"
        assert '台北市信義區' in text, "Nested company.address replacement failed"

        print("✓ All simple replacements working correctly!")
        print(f"\nOutput preview:\n{text}\n")


def test_advanced_jinja2():
    """Test advanced Jinja2 features (loops, conditionals)."""
    print("\n=== Test 2: Advanced Jinja2 Features ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "advanced_template.docx"
        output_path = Path(tmpdir) / "advanced_output.docx"

        # Create template
        create_advanced_template(template_path)

        # Prepare data with lists and conditionals
        data = {
            'name': '李四',
            'salary': 50000,  # > 40000, should trigger "高" condition
            'projects': [
                {'name': '專案A', 'budget': 100000},
                {'name': '專案B', 'budget': 200000},
                {'name': '專案C', 'budget': 150000},
            ]
        }

        # Process
        processor = DocxProcessor(template_path)
        processor.fill_template(data, output_path, verbose=True)

        # Verify
        doc = Document(output_path)
        text = '\n'.join([p.text for p in doc.paragraphs])

        assert '李四' in text, "Name replacement failed"
        assert '專案A: 100000元' in text, "Loop iteration 1 failed"
        assert '專案B: 200000元' in text, "Loop iteration 2 failed"
        assert '專案C: 150000元' in text, "Loop iteration 3 failed"
        assert '薪資等級：高' in text, "Conditional (if) failed"
        assert '薪資等級：一般' not in text, "Conditional (else) should not appear"

        print("✓ All advanced Jinja2 features working correctly!")
        print(f"\nOutput preview:\n{text}\n")


def test_placeholder_detection():
    """Test placeholder detection."""
    print("\n=== Test 3: Placeholder Detection ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = Path(tmpdir) / "detect_template.docx"

        # Create template
        create_simple_template(template_path)

        # Detect placeholders
        processor = DocxProcessor(template_path)
        placeholders = processor.get_placeholders()

        print(f"Found placeholders: {sorted(placeholders)}")

        expected = {'name', 'date', 'company.name', 'company.address'}
        assert placeholders == expected, f"Expected {expected}, got {placeholders}"

        print("✓ Placeholder detection working correctly!")


if __name__ == "__main__":
    print("Testing docxtpl integration...\n")

    try:
        test_simple_replacement()
        test_advanced_jinja2()
        test_placeholder_detection()

        print("\n" + "="*50)
        print("🎉 All tests passed!")
        print("="*50)
        print("\n✓ Simple {{key}} syntax works")
        print("✓ Nested {{obj.key}} syntax works")
        print("✓ Advanced {% for %} loops work")
        print("✓ Advanced {% if %} conditionals work")
        print("✓ Placeholder detection works")
        print("\n👉 No flag needed - all features work automatically!")

    except AssertionError as e:
        print(f"\n❌ Test failed: {e}")
        exit(1)
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        exit(1)
