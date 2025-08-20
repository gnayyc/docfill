#!/usr/bin/env python3
"""Test script to debug placeholder replacement issues."""

from docx import Document
from docx_processor import DocxProcessor
import re

def debug_paragraph_runs(paragraph):
    """Debug function to show how text is split across runs."""
    print(f"Full paragraph text: '{paragraph.text}'")
    print(f"Number of runs: {len(paragraph.runs)}")
    
    for i, run in enumerate(paragraph.runs):
        print(f"  Run {i}: '{run.text}'")
    
    # Check for placeholders
    placeholders = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
    print(f"Placeholders found: {placeholders}")
    print("-" * 50)

def test_with_simple_docx():
    """Create a simple test document."""
    doc = Document()
    
    # Add test paragraph with placeholder
    p1 = doc.add_paragraph("Name: {{name}}")
    p2 = doc.add_paragraph("Date: {{date}}")
    p3 = doc.add_paragraph("Description: {{description}}")
    
    # Save test document
    doc.save("test_template.docx")
    print("Created test_template.docx")
    
    # Debug the created document
    print("\nDebugging test_template.docx:")
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {i}:")
        debug_paragraph_runs(paragraph)

if __name__ == "__main__":
    test_with_simple_docx()