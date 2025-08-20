#!/usr/bin/env python3
"""Test formatting preservation with complex scenarios."""

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx_processor import DocxProcessor

def create_complex_test_doc():
    """Create a test document with various formatting."""
    doc = Document()
    
    # Test 1: Mixed formatting in one paragraph
    p1 = doc.add_paragraph()
    p1.add_run("Name: ").bold = True
    p1.add_run("{{name}}")
    p1.add_run(" (Employee)").italic = True
    
    # Test 2: Different font sizes
    p2 = doc.add_paragraph()
    run1 = p2.add_run("Date: ")
    run1.font.size = Pt(14)
    run1.bold = True
    
    run2 = p2.add_run("{{date}}")
    run2.font.size = Pt(12)
    
    # Test 3: Color formatting
    p3 = doc.add_paragraph()
    run3 = p3.add_run("Department: ")
    run3.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)  # Blue
    
    run4 = p3.add_run("{{department}}")
    run4.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red
    
    # Test 4: Simulate split placeholder (this might happen when editing in Word)
    p4 = doc.add_paragraph()
    p4.add_run("Description: {{desc")
    p4.add_run("ription}}")  # Split placeholder
    
    doc.save("complex_template.docx")
    print("Created complex_template.docx with various formatting")
    
    return doc

def debug_document_runs(doc_path):
    """Debug the runs structure of a document."""
    doc = Document(doc_path)
    print(f"\nDebugging: {doc_path}")
    print("=" * 50)
    
    for p_idx, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {p_idx}: '{paragraph.text}'")
        print(f"  Number of runs: {len(paragraph.runs)}")
        
        for r_idx, run in enumerate(paragraph.runs):
            formatting = []
            if run.bold:
                formatting.append("BOLD")
            if run.italic:
                formatting.append("ITALIC")
            if run.underline:
                formatting.append("UNDERLINE")
            if run.font.name:
                formatting.append(f"Font:{run.font.name}")
            if run.font.size:
                formatting.append(f"Size:{run.font.size.pt}")
            if run.font.color.rgb:
                formatting.append(f"Color:{run.font.color.rgb}")
                
            format_str = f" [{', '.join(formatting)}]" if formatting else ""
            print(f"    Run {r_idx}: '{run.text}'{format_str}")
        print()

if __name__ == "__main__":
    # Create test document
    create_complex_test_doc()
    
    # Debug original
    debug_document_runs("complex_template.docx")
    
    # Process with our tool
    print("\nProcessing with FilDoc...")
    processor = DocxProcessor("complex_template.docx")
    
    # Create simple config for testing
    test_data = {
        "name": "張三",
        "date": "2024-08-19", 
        "department": "資訊部",
        "description": "軟體工程師"
    }
    
    processor.fill_template(test_data, "complex_output.docx")
    
    # Debug output
    debug_document_runs("complex_output.docx")