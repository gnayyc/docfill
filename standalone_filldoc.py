#!/usr/bin/env python3
"""
Standalone FillDoc - DOCX Template Filler
Single file version with all dependencies inline.

Usage: python standalone_filldoc.py template.docx config.yaml
"""

import argparse
import sys
import re
import json
from pathlib import Path
from typing import Dict, Any, Union

try:
    import yaml
    from docx import Document
except ImportError:
    print("Error: Required packages not installed.")
    print("Please install: pip install python-docx PyYAML")
    sys.exit(1)

# Inline all the classes here for standalone use
class ConfigReader:
    """Reads configuration data from YAML and JSON files."""
    
    def __init__(self, config_path: Union[str, Path]):
        self.config_path = Path(config_path)
        if not self.config_path.exists():
            raise FileNotFoundError(f"Config file not found: {config_path}")
    
    def read(self) -> Dict[str, str]:
        file_extension = self.config_path.suffix.lower()
        
        if file_extension in ['.yaml', '.yml']:
            return self._read_yaml()
        elif file_extension == '.json':
            return self._read_json()
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _read_yaml(self) -> Dict[str, str]:
        with open(self.config_path, 'r', encoding='utf-8') as file:
            data = yaml.safe_load(file)
        return self._flatten_dict(data)
    
    def _read_json(self) -> Dict[str, str]:
        with open(self.config_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        return self._flatten_dict(data)
    
    def _flatten_dict(self, data: Dict[str, Any], parent_key: str = '', separator: str = '.') -> Dict[str, str]:
        items = []
        for key, value in data.items():
            new_key = f"{parent_key}{separator}{key}" if parent_key else key
            
            if isinstance(value, dict):
                items.extend(self._flatten_dict(value, new_key, separator).items())
            else:
                items.append((new_key, str(value)))
        
        return dict(items)

class DocxProcessor:
    """Process DOCX templates by replacing {{placeholder}} with actual values."""
    
    def __init__(self, template_path: Union[str, Path]):
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
    
    def fill_template(self, data: Dict[str, str], output_path: Union[str, Path]) -> None:
        doc = Document(self.template_path)
        
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data)
        
        doc.save(output_path)
    
    def _replace_in_paragraph(self, paragraph, data: Dict[str, str]) -> None:
        for run in paragraph.runs:
            original_text = run.text
            if not original_text:
                continue
            
            placeholders = re.findall(r'\{\{([^}]+)\}\}', original_text)
            
            if placeholders:
                new_text = original_text
                for placeholder in placeholders:
                    placeholder_with_braces = f"{{{{{placeholder}}}}}"
                    
                    if placeholder in data:
                        replacement = data[placeholder]
                    elif placeholder.lower() in [k.lower() for k in data.keys()]:
                        actual_key = next(k for k in data.keys() if k.lower() == placeholder.lower())
                        replacement = data[actual_key]
                    else:
                        replacement = placeholder_with_braces
                        print(f"Warning: No replacement found for placeholder: {placeholder}")
                    
                    new_text = new_text.replace(placeholder_with_braces, replacement)
                
                if new_text != original_text:
                    run.text = new_text

def main():
    parser = argparse.ArgumentParser(description="Fill DOCX templates with configuration data")
    parser.add_argument('template', help='Path to DOCX template file')
    parser.add_argument('config', help='Path to configuration file (YAML/JSON)')
    parser.add_argument('-o', '--output', help='Output file path')
    
    args = parser.parse_args()
    
    try:
        # Determine output path
        if args.output:
            output_path = Path(args.output)
        else:
            template_path = Path(args.template)
            stem = template_path.stem
            suffix = template_path.suffix
            output_path = template_path.parent / f"{stem}_filled{suffix}"
        
        # Process
        processor = DocxProcessor(args.template)
        reader = ConfigReader(args.config)
        config_data = reader.read()
        
        processor.fill_template(config_data, output_path)
        print(f"Template filled successfully: {output_path}")
        
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    
    return 0

if __name__ == "__main__":
    sys.exit(main())