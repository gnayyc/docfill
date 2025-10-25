"""DOCX template processor using docxtpl for replacing placeholders with values."""

import re
from pathlib import Path
from typing import Dict, Union
from docxtpl import DocxTemplate


class DocxProcessor:
    """Process DOCX templates by replacing {{placeholder}} with actual values.

    Supports both simple variable substitution and advanced Jinja2 features:
    - Simple: {{name}}, {{company.name}}
    - Advanced: {% for item in items %}, {% if condition %}

    No configuration needed - the template syntax determines what features are used.
    """

    def __init__(self, template_path: Union[str, Path]):
        """Initialize with template file path.

        Args:
            template_path: Path to the DOCX template file

        Raises:
            FileNotFoundError: If template file doesn't exist
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")

    def fill_template(self, data: Dict[str, str], output_path: Union[str, Path], verbose: bool = False) -> None:
        """Fill template with data and save to output path.

        Uses docxtpl which automatically handles:
        - Simple variable substitution: {{key}}
        - Nested data access: {{company.name}}
        - Jinja2 advanced features: loops, conditionals, filters

        Args:
            data: Dictionary of placeholder keys and replacement values
            output_path: Path where the filled document will be saved
            verbose: If True, print detailed progress information

        Raises:
            RuntimeError: If template processing or saving fails
        """
        try:
            if verbose:
                print(f"  Loading template: {self.template_path}")
            doc = DocxTemplate(self.template_path)
        except Exception as e:
            raise RuntimeError(f"Failed to open template '{self.template_path}': {e}")

        try:
            if verbose:
                print(f"  Rendering template with {len(data)} configuration items")

            # docxtpl's render() automatically handles all Jinja2 syntax
            # including simple {{key}}, nested {{obj.key}}, and advanced {% %} tags
            doc.render(data)

        except Exception as e:
            raise RuntimeError(f"Failed to render template: {e}")

        # Save the document
        try:
            if verbose:
                print(f"  Saving to: {output_path}")
            doc.save(output_path)
            if verbose:
                print(f"  Successfully saved: {output_path}")
        except Exception as e:
            raise RuntimeError(f"Failed to save document to '{output_path}': {e}")

    def get_placeholders(self) -> set:
        """Extract all placeholders from the template.

        Returns all {{variable}} style placeholders found in the document.
        Note: This only detects simple {{}} placeholders, not Jinja2 control structures.

        Returns:
            Set of placeholder names (without the {{ }} braces)
        """
        # We need to read the raw document XML to find placeholders
        # since docxtpl doesn't provide a direct method for this
        from docx import Document

        doc = Document(self.template_path)
        placeholders = set()

        # Pattern to match {{placeholder}} syntax
        # This will work for both simple and Jinja2 templates
        pattern = r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_.]*)\s*\}\}'

        # Get placeholders from paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)

        # Get placeholders from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)

        # Get placeholders from headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    placeholders.update(matches)

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    placeholders.update(matches)

        return placeholders
